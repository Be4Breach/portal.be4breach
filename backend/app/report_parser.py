import re
import os
import statistics
from collections import Counter, defaultdict
from tempfile import NamedTemporaryFile
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from base64 import b64encode


SEVERITY_ORDER = ["critical", "high", "medium", "low", "informational"]
SEVERITY_COLORS = {
    "critical": "hsl(0, 72%, 51%)",
    "high": "hsl(25, 95%, 53%)",
    "medium": "hsl(45, 93%, 47%)",
    "low": "hsl(0, 0%, 64%)",
    "informational": "hsl(210, 10%, 70%)",
}


def extract_text_from_docx(file_path: str) -> str:
    document = Document(file_path)
    paragraphs = [para.text for para in document.paragraphs]
    return "\n".join(paragraphs)


def extract_engagement_details(text: str, doc: Document) -> Dict[str, str]:
    """Pull client, report date, and audit type from paragraphs and the metadata table."""
    engagement: Dict[str, str] = {}

    client_match = re.search(r"Prepared For:\s*(.+)", text, flags=re.IGNORECASE)
    date_match = re.search(r"Report Date:\s*([0-9]{2}-[0-9]{2}-[0-9]{4})", text, flags=re.IGNORECASE)
    audit_type_match = re.search(r"Type of Audit[:\s]*([A-Za-z ]+)", text, flags=re.IGNORECASE)

    if client_match:
        engagement["client"] = client_match.group(1).strip()

    if date_match:
        engagement["reportDate"] = date_match.group(1).strip()

    if audit_type_match:
        engagement["auditType"] = audit_type_match.group(1).strip()

    # Fallback to first metadata table if present
    if doc.tables:
        meta_table = doc.tables[0]
        for row in meta_table.rows:
            if len(row.cells) < 2:
                continue
            key = row.cells[0].text.strip().lower()
            val = row.cells[1].text.strip()
            if not val:
                continue
            if "type of audit" in key and "auditType" not in engagement:
                engagement["auditType"] = val
            if "effective date" in key and "reportDate" not in engagement:
                engagement["reportDate"] = val
            if "document title" in key and "title" not in engagement:
                engagement["title"] = val

    return engagement


def _safe_float(value: str) -> Optional[float]:
    match = re.search(r"(\d+(?:\.\d+)?)", value or "")
    if match:
        try:
            return float(match.group(1))
        except ValueError:
            return None
    return None


def _clean_text(value: str) -> str:
    return (value or "").replace("\n", " ").replace("\r", " ").strip()


def _extract_images_from_cell(cell) -> List[str]:
    """
    Return list of data:// URIs for images embedded within a table cell.
    """
    images: List[str] = []
    # find all blip elements that reference embedded images
    # python-docx does not expose xpath with namespaces param on this object, so use .xml string search via lxml element
    for blip in cell._tc.xpath(".//a:blip"):
        r_id = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
        if not r_id:
            continue
        image_part = cell.part.related_parts.get(r_id)
        if not image_part:
            continue
        content_type = getattr(image_part, "content_type", "image/png")
        data_uri = f"data:{content_type};base64,{b64encode(image_part.blob).decode('ascii')}"
        images.append(data_uri)
    return images


def extract_findings_from_tables(doc: Document) -> List[Dict]:
    """
    Parse findings from the summary table and the per-finding detail tables.
    Expected format (as in Docon VAPT reports):
    - A summary table with header containing 'Observation/ Vulnerability title'
    - Subsequent detail tables where first cell starts with '<id>: <title>'
    """
    findings: Dict[int, Dict] = {}

    # 1) Summary table
    for table in doc.tables:
        header = [cell.text.lower() for cell in table.rows[0].cells]
        if any("observation" in h and "vulnerability" in h for h in header):
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) < 6:
                    continue
                id_raw = cells[0].strip()
                try:
                    fid = int(re.sub(r"\D", "", id_raw))
                except ValueError:
                    continue
                finding = findings.setdefault(fid, {"id": fid})
                if cells[2]:
                    finding["title"] = _clean_text(cells[2])
                if cells[3]:
                    finding["cwe"] = _clean_text(cells[3])
                if cells[4]:
                    cvss = _safe_float(cells[4])
                    if cvss is not None:
                        finding["cvssScore"] = cvss
                if cells[5]:
                    finding["severity"] = cells[5].strip().title()
            break  # assume only one summary table

    # 2) Detail tables (one per finding, structured as key/value rows)
    for table in doc.tables:
        first_cell = table.rows[0].cells[0].text if table.rows else ""
        id_match = re.match(r"(\d+):\s*(.+)", first_cell)
        if not id_match:
            continue
        fid = int(id_match.group(1))
        title = _clean_text(id_match.group(2))
        finding = findings.setdefault(fid, {"id": fid})
        if title:
            finding["title"] = title

        skip_next = False
        rows = table.rows
        for i in range(1, len(rows)):
            if skip_next:
                skip_next = False
                continue

            row = rows[i]
            if len(row.cells) < 2:
                continue

            key_raw = row.cells[0].text.strip()
            key = key_raw.lower()

            def _row_content(r) -> str:
                first = _clean_text(r.cells[0].text)
                second = _clean_text(r.cells[1].text)
                return first or second

            def next_row_text() -> Optional[str]:
                nxt_idx = i + 1
                if nxt_idx < len(rows):
                    nxt = rows[nxt_idx]
                    content = _row_content(nxt)
                    if content and content.lower() != key_raw.lower():
                        return content
                return None

            val = _clean_text(row.cells[1].text)

            if "severity" in key:
                finding["severity"] = val.title()
            elif key == "status":
                finding["status"] = val.title()
            elif "cve" in key or "cwe" in key:
                finding["cwe"] = val
            elif "cvss" in key:
                cv = _safe_float(val)
                if cv is not None:
                    if "cvssScore" not in finding or cv > (finding.get("cvssScore") or 0):
                        finding["cvssScore"] = cv
            elif "description" in key:
                content = next_row_text() or val
                finding["description"] = content
                if next_row_text():
                    skip_next = True
            elif key.startswith("impact"):
                content = next_row_text() or val
                finding["impact"] = content
                if next_row_text():
                    skip_next = True
            elif "affected asset" in key:
                content = next_row_text() or val
                finding["affectedAsset"] = content
                if next_row_text():
                    skip_next = True
            elif "recommendation" in key:
                content = next_row_text() or val
                finding["recommendations"] = content
                if next_row_text():
                    skip_next = True
            elif "reference" in key:
                content = next_row_text() or val
                finding["references"] = content
                if next_row_text():
                    skip_next = True
            elif "proof of concept" in key:
                content = next_row_text() or val
                finding["poc"] = content
                if next_row_text():
                    skip_next = True
                # collect images from current and next row if present
                poc_images = _extract_images_from_cell(row.cells[1])
                nxt_idx = i + 1
                if next_row_text() and nxt_idx < len(rows):
                    poc_images.extend(_extract_images_from_cell(rows[nxt_idx].cells[1]))
                if poc_images:
                    existing = finding.get("pocImages", [])
                    finding["pocImages"] = list({*existing, *poc_images})

    return list(findings.values())


def generate_summary(findings: List[Dict]) -> Dict[str, int]:
    summary: Dict[str, int] = defaultdict(int)

    for finding in findings:
        severity = (finding.get("severity") or "").lower()
        if severity in SEVERITY_ORDER:
            summary[severity] += 1

    return dict(summary)


def parse_report(file_path: str) -> Dict:
    doc = Document(file_path)
    text = "\n".join([p.text for p in doc.paragraphs])

    engagement = extract_engagement_details(text, doc)
    findings = extract_findings_from_tables(doc)
    summary = generate_summary(findings)

    return {
        "engagement": engagement,
        "summary": summary,
        "totalFindings": len(findings),
        "findings": findings,
    }


def _normalize_status(status: Optional[str]) -> str:
    if not status:
        return "Unknown"

    normalized = status.strip().lower()
    if "open" in normalized and "closed" not in normalized:
        return "Open"
    if "progress" in normalized or "wip" in normalized:
        return "In Progress"
    if "resolved" in normalized or "closed" in normalized or "fixed" in normalized:
        return "Resolved"
    if "accepted" in normalized:
        return "Accepted"
    return status.title()


def _severity_sort_key(severity: str) -> Tuple[int, str]:
    idx = SEVERITY_ORDER.index(severity.lower()) if severity and severity.lower() in SEVERITY_ORDER else len(SEVERITY_ORDER)
    return (idx, severity)


def build_report_dashboard(parsed: Dict) -> Dict:
    summary = parsed.get("summary", {})
    findings = parsed.get("findings", [])

    severity_chart = [
        {
            "name": level.title(),
            "value": summary.get(level, 0),
            "color": SEVERITY_COLORS[level],
        }
        for level in SEVERITY_ORDER
    ]

    status_counts = Counter(_normalize_status(f.get("status")) for f in findings)
    status_breakdown = [{"status": k, "count": v} for k, v in status_counts.items() if v]

    scores = [f["cvssScore"] for f in findings if isinstance(f.get("cvssScore"), (int, float))]
    cvss = None
    if scores:
        cvss = {
            "average": round(statistics.mean(scores), 1),
            "max": max(scores),
            "min": min(scores),
            "count": len(scores),
        }

    sorted_findings = sorted(
        findings,
        key=lambda f: (
            _severity_sort_key(f.get("severity", "")),
            -(f.get("cvssScore", 0) or 0),
        ),
    )

    return {
        **parsed,
        "severityChart": severity_chart,
        "statusBreakdown": status_breakdown,
        "cvss": cvss,
        "topFindings": sorted_findings[:5],
    }


def parse_uploaded_file(upload_file) -> Dict:
    with NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        content = upload_file.file.read()
        tmp.write(content)
        tmp_path = tmp.name

    try:
        parsed = parse_report(tmp_path)
        return build_report_dashboard(parsed)
    finally:
        try:
            os.remove(tmp_path)
        except OSError:
            pass
